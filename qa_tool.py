"""
QA Screenshot Capture Tool
Portable Windows application — no admin or installation required.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import threading
import ctypes
import ctypes.wintypes as wintypes
import os
import sys
import json
import time
from datetime import datetime
from io import BytesIO

try:
    from PIL import ImageGrab
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as exc:
    import tkinter.messagebox as _mb
    _mb.showerror(
        "Missing dependencies",
        f"Required package missing: {exc}\n\n"
        "pip install Pillow python-docx",
    )
    sys.exit(1)


# ── Hotkey manager (Windows RegisterHotKey API) ───────────────────────────────
# Uses the official Windows hotkey API instead of a keyboard hook library.
# RegisterHotKey does NOT install a system-wide hook, so antivirus software
# has no reason to flag it — it's the same mechanism Word, Snipping Tool, etc. use.

class HotkeyManager:
    WM_HOTKEY = 0x0312
    WM_QUIT   = 0x0012
    _ID       = 1

    _VK = {
        'f1': 0x70, 'f2': 0x71, 'f3': 0x72,  'f4': 0x73,
        'f5': 0x74, 'f6': 0x75, 'f7': 0x76,  'f8': 0x77,
        'f9': 0x78, 'f10':0x79, 'f11':0x7A,  'f12':0x7B,
        'print screen': 0x2C, 'snapshot': 0x2C,
    }
    _MOD = {'ctrl': 0x0002, 'shift': 0x0004, 'alt': 0x0001}

    def __init__(self):
        self._u32      = ctypes.windll.user32
        self._k32      = ctypes.windll.kernel32
        self._callback = None
        self._thread   = None
        self._win_tid  = None
        self._stop     = threading.Event()

    def _parse(self, hotkey: str):
        parts = [p.strip().lower() for p in hotkey.split('+')]
        key   = parts[-1]
        mod_v = 0
        for m in parts[:-1]:
            mod_v |= self._MOD.get(m, 0)
        vk = self._VK.get(key)
        if vk is None and len(key) == 1:
            vk = ord(key.upper())
        return mod_v, vk or 0

    def register(self, hotkey: str, callback):
        self.unregister()
        self._stop.clear()
        self._callback = callback
        mod_v, vk = self._parse(hotkey)
        self._thread = threading.Thread(
            target=self._loop, args=(mod_v, vk), daemon=True, name="hotkey")
        self._thread.start()

    def unregister(self):
        if self._thread and self._thread.is_alive():
            self._stop.set()
            if self._win_tid:
                self._u32.PostThreadMessageW(self._win_tid, self.WM_QUIT, 0, 0)
            self._thread.join(timeout=2)
        self._thread   = None
        self._callback = None
        self._win_tid  = None

    def _loop(self, mod_v: int, vk: int):
        self._win_tid = self._k32.GetCurrentThreadId()
        if not self._u32.RegisterHotKey(None, self._ID, mod_v, vk):
            return
        msg = wintypes.MSG()
        while not self._stop.is_set():
            ret = self._u32.GetMessageW(ctypes.byref(msg), None, 0, 0)
            if ret <= 0:
                break
            if msg.message == self.WM_HOTKEY and msg.wParam == self._ID:
                cb = self._callback
                if cb:
                    cb()
        self._u32.UnregisterHotKey(None, self._ID)


def app_dir() -> str:
    # When launched via QAScreenshotToolPortable.exe the launcher sets this
    # so config.json lands in Data\settings\ (portable root) not beside the exe
    portable = os.environ.get("PORTABLE_DATA_DIR")
    if portable:
        os.makedirs(portable, exist_ok=True)
        return portable
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _icon_path() -> str:
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, "icon.ico")
    return os.path.join(app_dir(), "icon.ico")


# ── Splash screen ─────────────────────────────────────────────────────────────

class SplashScreen:
    """
    Full-screen-style splash that shows for 3 seconds then reveals the app.
    Message: 'Built for Carrots.. Always be grateful to your brother, Sistaaa'
    """
    DURATION_MS = 3000

    def __init__(self, root: tk.Tk):
        self._root = root

        dlg = tk.Toplevel(root)
        dlg.overrideredirect(True)          # no title bar / border
        dlg.attributes("-topmost", True)

        W, H = 520, 340
        sw = dlg.winfo_screenwidth()
        sh = dlg.winfo_screenheight()
        dlg.geometry(f"{W}x{H}+{(sw - W) // 2}+{(sh - H) // 2}")

        BG   = "#4A154B"   # deep plum
        ACC  = "#7C3085"   # lighter purple border
        PINK = "#FF79C6"   # hot pink accent

        # Outer border frame
        border = tk.Frame(dlg, bg=ACC)
        border.place(x=0, y=0, width=W, height=H)

        # Inner background
        inner = tk.Frame(border, bg=BG)
        inner.place(x=2, y=2, width=W - 4, height=H - 4)

        # Scissors emoji
        tk.Label(inner, text="✂", font=("Segoe UI Emoji", 54),
                 bg=BG, fg="white").pack(pady=(28, 4))

        # App name
        tk.Label(inner, text="QA Screenshot Tool",
                 font=("Segoe UI", 13, "bold"),
                 bg=BG, fg="#DDA0DD").pack()          # plum text

        # Divider line
        tk.Frame(inner, bg=ACC, height=1).pack(fill="x", padx=60, pady=10)

        # "Built for Carrots.."
        tk.Label(inner,
                 text="Built for Carrots..",
                 font=("Segoe UI", 16, "bold"),
                 bg=BG, fg=PINK).pack(pady=(0, 4))

        # Sister message
        tk.Label(inner,
                 text="Always be grateful to your brother, Sistaaa",
                 font=("Segoe UI", 11, "italic"),
                 bg=BG, fg="#F8C8F0",
                 wraplength=420, justify="center").pack()

        # Progress bar
        tk.Frame(inner, bg=BG, height=12).pack()
        style = ttk.Style(dlg)
        style.configure(
            "Splash.Horizontal.TProgressbar",
            troughcolor="#2D0A30",
            background=PINK,
            bordercolor=BG,
            lightcolor=PINK,
            darkcolor=PINK,
        )
        self._bar = ttk.Progressbar(
            inner, length=340, mode="determinate",
            style="Splash.Horizontal.TProgressbar",
        )
        self._bar.pack(pady=(6, 4))

        tk.Label(inner, text="Click anywhere to skip",
                 font=("Segoe UI", 7), bg=BG, fg="#7C3085").pack()

        dlg.bind("<Button-1>", lambda _: self._finish())
        self._dlg   = dlg
        self._steps = 0
        self._interval = 40                         # ms per tick
        self._total_ticks = self.DURATION_MS // self._interval
        self._animate()

    def _animate(self):
        self._steps += 1
        self._bar["value"] = min(100, self._steps * 100 // self._total_ticks)
        if self._steps < self._total_ticks:
            self._dlg.after(self._interval, self._animate)
        else:
            self._finish()

    def _finish(self):
        try:
            self._dlg.destroy()
        except Exception:
            pass
        try:
            self._root.deiconify()
        except Exception:
            pass


# ── Settings ──────────────────────────────────────────────────────────────────

class Settings:
    _DEFAULTS = {
        "hotkey": "f1",
        "auto_timestamp": True,
        "auto_number": True,
        "prompt_description": False,
        "minimize_on_capture": True,
        "capture_delay": 0,
        "tester_name": "",
        "project_name": "",
        "save_path": "",
    }

    def __init__(self):
        self._path = os.path.join(app_dir(), "config.json")
        self._data: dict = dict(self._DEFAULTS)
        self._load()

    def _load(self):
        try:
            if os.path.exists(self._path):
                with open(self._path, "r") as fh:
                    self._data.update(json.load(fh))
        except Exception:
            pass

    def save(self):
        try:
            with open(self._path, "w") as fh:
                json.dump(self._data, fh, indent=2)
        except Exception:
            pass

    def __getitem__(self, key):
        return self._data.get(key, self._DEFAULTS.get(key))

    def __setitem__(self, key, value):
        self._data[key] = value


# ── Main application ──────────────────────────────────────────────────────────

class QATool:
    APP_TITLE = "QA Screenshot Tool"
    WIN_SIZE  = "540x500"

    def __init__(self):
        self.root = tk.Tk()
        self.root.withdraw()                   # hide until splash finishes
        self.root.title(self.APP_TITLE)
        self.root.geometry(self.WIN_SIZE)
        self.root.resizable(False, False)

        # Window icon (title bar + taskbar)
        ip = _icon_path()
        if os.path.exists(ip):
            try:
                self.root.iconbitmap(ip)
            except Exception:
                pass

        self.settings = Settings()

        self.active     = False
        self.document   = None
        self.doc_path   = ""
        self.count      = 0
        self.sess_start = None
        self._lock      = threading.Lock()
        self._hotkey    = HotkeyManager()

        self._build_ui()
        self._restore_settings_to_ui()

        # Show splash — main window appears automatically when it finishes
        SplashScreen(self.root)

    # ── UI ───────────────────────────────────────────────────────────────────

    def _build_ui(self):
        nb = ttk.Notebook(self.root)
        nb.pack(fill="both", expand=True, padx=8, pady=8)

        capture_tab  = ttk.Frame(nb, padding=12)
        settings_tab = ttk.Frame(nb, padding=12)
        about_tab    = ttk.Frame(nb, padding=12)

        nb.add(capture_tab,  text="  Capture  ")
        nb.add(settings_tab, text="  Settings  ")
        nb.add(about_tab,    text="  About  ")

        self._build_capture_tab(capture_tab)
        self._build_settings_tab(settings_tab)
        self._build_about_tab(about_tab)

    def _build_capture_tab(self, p):
        tk.Label(p, text="QA Screenshot Capture Tool",
                 font=("Segoe UI", 13, "bold")).pack(pady=(4, 2))
        tk.Label(p, text="Capture test screenshots directly into Word documents",
                 font=("Segoe UI", 9), fg="#555").pack(pady=(0, 10))

        doc = ttk.LabelFrame(p, text="Document", padding=8)
        doc.pack(fill="x", pady=4)

        row = tk.Frame(doc)
        row.pack(fill="x", pady=2)
        tk.Label(row, text="File name:", width=12, anchor="w").pack(side="left")
        self.fname_var = tk.StringVar(
            value=f"TestReport_{datetime.now().strftime('%Y%m%d')}")
        ttk.Entry(row, textvariable=self.fname_var, width=34).pack(side="left", padx=4)

        row2 = tk.Frame(doc)
        row2.pack(fill="x", pady=2)
        tk.Label(row2, text="Save to:", width=12, anchor="w").pack(side="left")
        self.save_path_var = tk.StringVar(value=os.path.expanduser("~/Desktop"))
        ttk.Entry(row2, textvariable=self.save_path_var, width=28).pack(side="left", padx=4)
        ttk.Button(row2, text="Browse", command=self._browse, width=7).pack(side="left")

        hk_frame = ttk.LabelFrame(p, text="Active hotkey", padding=6)
        hk_frame.pack(fill="x", pady=4)
        self.hotkey_hint_var = tk.StringVar()
        tk.Label(hk_frame, textvariable=self.hotkey_hint_var,
                 font=("Segoe UI", 10, "bold"), fg="#006400").pack()

        btn = tk.Frame(p)
        btn.pack(pady=10)
        self.start_btn = ttk.Button(btn, text="Start Session",
                                    command=self.start_session, width=20)
        self.start_btn.pack(side="left", padx=6)
        self.stop_btn  = ttk.Button(btn, text="Stop & Save",
                                    command=self.stop_session,
                                    state="disabled", width=18)
        self.stop_btn.pack(side="left", padx=6)

        sf = ttk.LabelFrame(p, text="Session status", padding=8)
        sf.pack(fill="x", pady=4)
        self.status_var = tk.StringVar(
            value="No active session. Enter a file name and press Start.")
        tk.Label(sf, textvariable=self.status_var, font=("Segoe UI", 9),
                 fg="#003399", wraplength=460).pack()
        self.count_var = tk.StringVar(value="")
        tk.Label(sf, textvariable=self.count_var,
                 font=("Segoe UI", 9, "bold")).pack()
        self.last_var = tk.StringVar(value="")
        tk.Label(sf, textvariable=self.last_var,
                 font=("Segoe UI", 8), fg="#666").pack()

    def _build_settings_tab(self, p):
        hk = ttk.LabelFrame(p, text="Hotkey", padding=8)
        hk.pack(fill="x", pady=6)

        r = tk.Frame(hk)
        r.pack(fill="x")
        tk.Label(r, text="Screenshot key:", width=18, anchor="w").pack(side="left")
        self.hotkey_var = tk.StringVar(value="f1")
        hotkeys = [
            "f1", "f2", "f3", "f4", "f5", "f6",
            "f7", "f8", "f9", "f10", "f11", "f12",
            "ctrl+shift+s", "ctrl+shift+c",
            "ctrl+alt+s",   "ctrl+alt+c",
            "print screen",
        ]
        ttk.Combobox(r, textvariable=self.hotkey_var,
                     values=hotkeys, width=20).pack(side="left", padx=4)
        ttk.Button(r, text="Apply", command=self._apply_hotkey,
                   width=8).pack(side="left", padx=4)
        tk.Label(hk, text="Takes effect immediately in an active session.",
                 font=("Segoe UI", 8), fg="#777").pack(anchor="w", pady=(4, 0))

        co = ttk.LabelFrame(p, text="Capture options", padding=8)
        co.pack(fill="x", pady=6)

        self.auto_ts_var  = tk.BooleanVar(value=True)
        self.auto_num_var = tk.BooleanVar(value=True)
        self.prompt_d_var = tk.BooleanVar(value=False)
        self.minimize_var = tk.BooleanVar(value=True)
        self.delay_var    = tk.StringVar(value="0")

        ttk.Checkbutton(co, text="Add timestamp below each screenshot",
                        variable=self.auto_ts_var).pack(anchor="w", pady=2)
        ttk.Checkbutton(co, text="Auto-number screenshots (Screenshot 1, 2, 3 ...)",
                        variable=self.auto_num_var).pack(anchor="w", pady=2)
        ttk.Checkbutton(co, text="Prompt for description / test-step before each capture",
                        variable=self.prompt_d_var).pack(anchor="w", pady=2)
        ttk.Checkbutton(co, text="Minimize window during capture (keeps screenshots clean)",
                        variable=self.minimize_var).pack(anchor="w", pady=2)

        dr = tk.Frame(co)
        dr.pack(anchor="w", pady=2)
        tk.Label(dr, text="Capture delay (0-5 sec):").pack(side="left")
        ttk.Spinbox(dr, from_=0, to=5, textvariable=self.delay_var,
                    width=5).pack(side="left", padx=6)

        dh = ttk.LabelFrame(p, text="Document header info", padding=8)
        dh.pack(fill="x", pady=6)

        r1 = tk.Frame(dh)
        r1.pack(fill="x", pady=2)
        tk.Label(r1, text="Tester name:", width=15, anchor="w").pack(side="left")
        self.tester_var = tk.StringVar()
        ttk.Entry(r1, textvariable=self.tester_var, width=28).pack(side="left", padx=4)

        r2 = tk.Frame(dh)
        r2.pack(fill="x", pady=2)
        tk.Label(r2, text="Project / App:", width=15, anchor="w").pack(side="left")
        self.project_var = tk.StringVar()
        ttk.Entry(r2, textvariable=self.project_var, width=28).pack(side="left", padx=4)

        ttk.Button(p, text="Save settings", command=self._save_settings,
                   width=16).pack(pady=8)

    def _build_about_tab(self, p):
        tk.Label(p, text="QA Screenshot Capture Tool",
                 font=("Segoe UI", 14, "bold")).pack(pady=(20, 4))
        tk.Label(p, text="Version 1.0  -  Portable, no installation required",
                 font=("Segoe UI", 9), fg="#555").pack()
        instructions = (
            "\n"
            "How to use\n"
            "----------\n"
            "1.  Enter a file name on the Capture tab.\n"
            "2.  Choose where to save the document.\n"
            "3.  Press  Start Session.\n"
            "4.  Switch to your application under test.\n"
            "5.  Press your hotkey (default F1) to capture a screenshot.\n"
            "    The tool minimizes briefly so it stays out of the screenshot.\n"
            "6.  Press  Stop & Save  when done.\n"
            "    The Word document opens automatically.\n\n"
            "Tips\n"
            "----\n"
            "* Change the hotkey on the Settings tab at any time.\n"
            "* Enable 'Prompt for description' to annotate each screenshot.\n"
            "* The config.json file next to the .exe stores your settings.\n"
            "* The exe is fully portable - copy it to a USB drive and run it anywhere."
        )
        tk.Label(p, text=instructions, font=("Segoe UI", 9), justify="left",
                 wraplength=440).pack(padx=10)

    # ── Settings ─────────────────────────────────────────────────────────────

    def _restore_settings_to_ui(self):
        s = self.settings
        self.tester_var.set(s["tester_name"])
        self.project_var.set(s["project_name"])
        self.hotkey_var.set(s["hotkey"])
        self.auto_ts_var.set(s["auto_timestamp"])
        self.auto_num_var.set(s["auto_number"])
        self.prompt_d_var.set(s["prompt_description"])
        self.minimize_var.set(s["minimize_on_capture"])
        self.delay_var.set(str(s["capture_delay"]))
        if s["save_path"]:
            self.save_path_var.set(s["save_path"])
        self._refresh_hotkey_hint()

    def _refresh_hotkey_hint(self):
        hk = self.settings["hotkey"].upper()
        self.hotkey_hint_var.set(f"Press  [ {hk} ]  to capture a screenshot")

    def _apply_hotkey(self):
        hk = self.hotkey_var.get().strip()
        if not hk:
            messagebox.showwarning("Invalid hotkey", "Hotkey cannot be empty.")
            return
        self.settings["hotkey"] = hk
        self.settings.save()
        if self.active:
            self._register_hotkey()
        self._refresh_hotkey_hint()
        messagebox.showinfo("Hotkey updated", f"Hotkey set to: {hk.upper()}")

    def _save_settings(self):
        s = self.settings
        s["hotkey"]              = self.hotkey_var.get().strip()
        s["auto_timestamp"]      = self.auto_ts_var.get()
        s["auto_number"]         = self.auto_num_var.get()
        s["prompt_description"]  = self.prompt_d_var.get()
        s["minimize_on_capture"] = self.minimize_var.get()
        s["capture_delay"]       = int(self.delay_var.get() or 0)
        s["tester_name"]         = self.tester_var.get().strip()
        s["project_name"]        = self.project_var.get().strip()
        s["save_path"]           = self.save_path_var.get().strip()
        s.save()
        self._refresh_hotkey_hint()
        messagebox.showinfo("Saved", "Settings saved successfully.")

    def _browse(self):
        folder = filedialog.askdirectory(initialdir=self.save_path_var.get())
        if folder:
            self.save_path_var.set(folder)

    # ── Session ───────────────────────────────────────────────────────────────

    def start_session(self):
        fname = self.fname_var.get().strip()
        if not fname:
            messagebox.showerror("Missing name", "Please enter a document file name.")
            return

        for ch in r'<>:"/\|?*':
            fname = fname.replace(ch, "_")
        if not fname.lower().endswith(".docx"):
            fname += ".docx"

        save_dir = self.save_path_var.get().strip() or os.path.expanduser("~/Desktop")
        os.makedirs(save_dir, exist_ok=True)

        self.doc_path = os.path.join(save_dir, fname)
        try:
            self._create_document()
        except Exception as exc:
            messagebox.showerror("Document error",
                                 f"Could not create document:\n{exc}")
            return

        self.count      = 0
        self.sess_start = datetime.now()
        self.active     = True

        self._register_hotkey()

        self.start_btn.config(state="disabled")
        self.stop_btn.config(state="normal")
        hk = self.settings["hotkey"].upper()
        self.status_var.set(f"Session active - press [{hk}] to capture a screenshot!")
        self.count_var.set("Screenshots captured: 0")
        self.last_var.set("")

    def stop_session(self):
        if not self.active:
            return
        self.active = False

        self._hotkey.unregister()

        if self.doc_path and os.path.exists(self.doc_path):
            try:
                self._append_summary()
            except Exception:
                pass

        self.start_btn.config(state="normal")
        self.stop_btn.config(state="disabled")
        self.status_var.set(f"Session ended - {self.count} screenshot(s) saved.")
        self.count_var.set("")
        self.last_var.set("")

        if self.doc_path and os.path.exists(self.doc_path):
            if messagebox.askyesno(
                    "Session complete",
                    f"{self.count} screenshot(s) captured.\n\n"
                    f"Saved to:\n{self.doc_path}\n\nOpen document now?"):
                try:
                    os.startfile(self.doc_path)
                except Exception:
                    messagebox.showinfo(
                        "File saved",
                        f"Document saved to:\n{self.doc_path}\n\n"
                        "Open it manually in Microsoft Word.")

    # ── Document ──────────────────────────────────────────────────────────────

    def _create_document(self):
        doc = Document()

        title = doc.add_heading("Test Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"

        def fill_row(i, label, value):
            cells = tbl.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True

        fill_row(0, "Tester",  self.tester_var.get()  or "-")
        fill_row(1, "Project", self.project_var.get() or "-")
        fill_row(2, "Date",    datetime.now().strftime("%Y-%m-%d"))
        fill_row(3, "File",    self.fname_var.get())

        doc.add_paragraph("")
        doc.add_heading("Test Screenshots", level=1)
        doc.save(self.doc_path)

    def _append_summary(self):
        # Reload from disk so any tester edits in Word are preserved
        doc = Document(self.doc_path)
        end_time = datetime.now()
        doc.add_heading("Session Summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Total screenshots: ").bold = True
        p.add_run(str(self.count))
        p.add_run("    ")
        p.add_run("Start: ").bold = True
        p.add_run(self.sess_start.strftime("%Y-%m-%d %H:%M:%S") if self.sess_start else "-")
        p.add_run("    ")
        p.add_run("End: ").bold = True
        p.add_run(end_time.strftime("%Y-%m-%d %H:%M:%S"))
        doc.save(self.doc_path)

    # ── Hotkey & capture ──────────────────────────────────────────────────────

    def _register_hotkey(self):
        try:
            self._hotkey.register(self.settings["hotkey"], self._on_hotkey)
        except Exception as exc:
            messagebox.showerror(
                "Hotkey error",
                f"Could not register hotkey '{self.settings['hotkey']}'.\n\n"
                f"Error: {exc}\n\n"
                "Try a different key in Settings.",
            )

    def _on_hotkey(self):
        if not self.active:
            return
        opts = {
            "delay":       self.settings["capture_delay"],
            "minimize":    self.settings["minimize_on_capture"],
            "auto_ts":     self.settings["auto_timestamp"],
            "auto_num":    self.settings["auto_number"],
            "prompt_desc": self.settings["prompt_description"],
        }
        threading.Thread(target=self._capture_flow, args=(opts,), daemon=True).start()

    def _capture_flow(self, opts: dict):
        if opts["prompt_desc"]:
            self.root.after(0, lambda: self._desc_dialog(opts))
        else:
            self._do_capture("", opts)

    def _desc_dialog(self, opts: dict):
        dlg = tk.Toplevel(self.root)
        dlg.title("Screenshot description")
        dlg.geometry("400x140")
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.lift()
        dlg.focus_force()
        dlg.attributes("-topmost", True)

        tk.Label(dlg, text="Description / test step (press Enter or leave blank):",
                 font=("Segoe UI", 9)).pack(pady=(12, 4))
        desc_var = tk.StringVar()
        entry = ttk.Entry(dlg, textvariable=desc_var, width=46)
        entry.pack(padx=14)
        entry.focus()

        btns = tk.Frame(dlg)
        btns.pack(pady=10)

        def capture_with_desc():
            desc = desc_var.get().strip()
            dlg.destroy()
            threading.Thread(target=self._do_capture, args=(desc, opts), daemon=True).start()

        def skip():
            dlg.destroy()
            threading.Thread(target=self._do_capture, args=("", opts), daemon=True).start()

        ttk.Button(btns, text="Capture", command=capture_with_desc, width=12).pack(side="left", padx=8)
        ttk.Button(btns, text="Skip",    command=skip,               width=8 ).pack(side="left", padx=4)
        entry.bind("<Return>", lambda _: capture_with_desc())

    def _do_capture(self, description: str, opts: dict):
        if not self._lock.acquire(blocking=False):
            return

        try:
            if opts["minimize"]:
                self.root.after(0, self.root.iconify)

            time.sleep(max(0.4, opts["delay"]))

            screenshot = ImageGrab.grab(all_screens=True)
            buf = BytesIO()
            screenshot.save(buf, format="PNG")
            buf.seek(0)

            self.count += 1
            n  = self.count
            ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # Reload from disk every time — preserves any reordering the tester
            # did in Word between captures, instead of overwriting with stale RAM copy
            doc = Document(self.doc_path)

            heading = f"Screenshot {n}" if opts["auto_num"] else "Screenshot"
            doc.add_heading(heading, level=2)

            if opts["auto_ts"]:
                para = doc.add_paragraph(f"Captured: {ts}")
                run  = para.runs[0]
                run.font.size      = Pt(8)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            if description:
                notes = doc.add_paragraph(f"Notes: {description}")
                notes.runs[0].bold = True

            doc.add_picture(buf, width=Inches(6.0))
            doc.add_paragraph("")
            doc.save(self.doc_path)

            self.root.after(0, lambda: self._update_status(n, ts))
            self.root.after(0, self.root.deiconify)

        except Exception as exc:
            self.root.after(0, self.root.deiconify)
            self.root.after(0, lambda: messagebox.showerror("Capture failed", str(exc)))
        finally:
            self._lock.release()

    def _update_status(self, n: int, ts: str):
        self.count_var.set(f"Screenshots captured: {n}")
        self.last_var.set(f"Last capture: {ts}")
        self.status_var.set(f"Screenshot {n} saved to document.")

    # ── Close ─────────────────────────────────────────────────────────────────

    def _on_close(self):
        if self.active:
            if not messagebox.askyesno(
                    "Quit", "A session is active. Stop and save before quitting?"):
                return
            self.stop_session()
        self._hotkey.unregister()
        self.root.destroy()

    def run(self):
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.mainloop()


if __name__ == "__main__":
    QATool().run()
